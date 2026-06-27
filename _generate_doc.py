"""Generate Model Tour documentation as a Word .docx file."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

# ── Helper: shade a table cell ─────────────────────────────────────────
def shade(cell, color='D9E2F3'):
    """Apply background shading to a table cell."""
    from docx.oxml.ns import qn
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)

def set_cell_width(cell, width):
    """Set explicit cell width."""
    from docx.oxml.ns import qn
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.makeelement(qn('w:tcW'), {
        qn('w:w'): str(width),
        qn('w:type'): 'dxa',
    })
    tcPr.append(tcW)

# ── Title Page ──────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Model Tour Module')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Complete Project Documentation')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('\nProduct: HyperMesh / HyperWorks Desktop\n').font.size = Pt(11)
meta.add_run('Version: model_tour branch\n').font.size = Pt(11)
meta.add_run('Last Updated: 2026-06-26\n').font.size = Pt(11)

doc.add_page_break()

# ── Table of Contents placeholder ───────────────────────────────────────
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Project Overview',
    '2. Architecture — Two-Layer Design',
    '3. Key Classes & Relationships',
    '    3.1 Core: tour.py',
    '    3.2 Core: dialog.py',
    '    3.3 HM Layer: hmmodeltour.py',
    '    3.4 HM Layer: hm_utils.py',
    '4. Complete Workflow',
    '    4.1 Start → Tour Running',
    '    4.2 Navigation Flow',
    '    4.3 Event Handling',
    '    4.4 Image Capture Workflow',
    '5. Template System',
    '6. File Inventory & Summary',
    '7. Design Patterns',
    '8. Recent Changes',
]
for item in toc_items:
    p = doc.add_paragraph(f'    {item}' if item.startswith('    ') else item)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if not item.startswith('    '):
        p.runs[0].bold = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 1
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('1. Project Overview', level=1)
doc.add_paragraph(
    'The Model Tour feature provides an interactive entity-navigation experience within '
    'HyperMesh. Users select a collection of model entities (Components, Parts, Connectors, '
    'etc.) and step through them with rendered HTML summaries, image capture capabilities, '
    'editable Jinja2 templates, and automatic graphics-area isolation.'
)
p = doc.add_paragraph()
p.add_run('Key Capabilities:').bold = True
bullets = [
    'Interactive navigation (next/previous) through entity collections with wrap-around',
    'Jinja2 + Markdown templating for entity-specific content rendering',
    'Live template editing with syntax highlighting and validation',
    'Graphics window capture to clipboard or file (PNG)',
    'Dual UI modes: full dock window (TourDialog) or compact toolbar (TourToolbar)',
    'Entity lifecycle event handling (create, delete, update, UID change)',
    'Singleton controller preserves state across UI mode switches',
    'Settings persistence via application XML settings tree',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

# ══════════════════════════════════════════════════════════════════════════
# SECTION 2
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('2. Architecture — Two-Layer Design', level=1)
doc.add_paragraph(
    'The module is organized into two distinct layers: a generic core layer and a '
    'HyperMesh-specific specialization layer. This separation allows the core tour engine '
    'to be reused with different host applications.'
)

table = doc.add_table(rows=3, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Layer', 'Directory', 'Files']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    shade(cell)

data = [
    ['Core', 'unity/plugins/hwd/scripts/\npython/modeltour/', 'tour.py, dialog.py, utils.py'],
    ['HM-Specific', 'hm/scripts/modeltour/', 'hmmodeltour.py, hm_utils.py'],
]
for r_idx, row_data in enumerate(data):
    for c_idx, val in enumerate(row_data):
        table.rows[r_idx + 1].cells[c_idx].text = val

doc.add_paragraph()
doc.add_heading('Core Layer', level=2)
doc.add_paragraph(
    'tour.py — Abstract base classes (ModelTour), Jinja2 template engine '
    '(_Template, _TemplateManager), entity navigation, event handling, image capture. '
    'dialog.py — UI framework: TourController (singleton mode-switcher), TourDialog '
    '(full dock window), TourToolbar (compact toolbar), settings persistence, '
    'inline template editing with syntax highlighting. '
    'utils.py — Base template helper functions (capture_image, eval_tcl, etc.).'
)

doc.add_heading('HM Layer', level=2)
doc.add_paragraph(
    'hmmodeltour.py — HMModelTour extends ModelTour with HyperMesh-specific signal '
    'handling (model lifecycle, element renumbering). Four entity handler classes: '
    'HMDefaultHandler, HMComponentHandler, HMPartHandler, HMConnectorHandler. '
    'Panel-based entity selection via Tcl. '
    'hm_utils.py — HM template utilities (mass, element details, entity fetching).'
)

# ══════════════════════════════════════════════════════════════════════════
# SECTION 3
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('3. Key Classes & Relationships', level=1)

# 3.1
doc.add_heading('3.1 Core: tour.py', level=2)

classes_tour = [
    ('_Template', 'Wraps a compiled Jinja2 template with entity source. '
     'render(output_format="text"|"html", **context) → renders via Jinja2, '
     'optionally converts Markdown → HTML.'),
    ('_TemplateManager (Singleton)', 'Owns all templates per entity type name. '
     'Loading chain: (1) IEntityHandler defaults → (2) User overrides from '
     'modeltour_templates.j2. Supports register, update, validate, reset, and '
     'persistence to settings directory.'),
    ('IEntityHandler', 'Base handler for entity isolation. Subclasses declare '
     'entity_type= and tour= in class statement for auto-registration. '
     'Methods: handle(), template(), tcl_name().'),
    ('ModelTour (Abstract)', 'Core tour engine. Holds Collection + current entity. '
     'Navigation (next/previous with wrap). Observer pattern for events '
     '(EVENT_CURRENT_CHANGED, EVENT_ENTITY_DELETED, EVENT_ENTITY_UPDATED, etc.). '
     'Per-class handler map and template context registry. Image capture via hwi + PIL. '
     'Signal suppression guard for re-entrant event prevention.'),
]

table = doc.add_table(rows=len(classes_tour) + 1, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Class', 'Description']):
    cell = table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    shade(cell)
for r_idx, (name, desc) in enumerate(classes_tour):
    table.rows[r_idx + 1].cells[0].text = name
    table.rows[r_idx + 1].cells[0].paragraphs[0].runs[0].bold = True
    table.rows[r_idx + 1].cells[1].text = desc

# 3.2
doc.add_heading('3.2 Core: dialog.py', level=2)

classes_dialog = [
    ('_TourSettings (Singleton)', 'Wraps hwx.xmlui.settings under '
     'Unity/DesktopUI/ModelTour/. Keys: capture_destination, capture_directory, '
     'capture_file_mode, include_adjacent, selection_entitytype, ui_mode.'),
    ('TourController (Singleton)', 'Owns the shared ModelTour instance. '
     'Manages two mutually exclusive UI modes: MODE_DIALOG and MODE_TOOLBAR. '
     'start() → select entities → switch mode. switch_mode() preserves tour state.'),
    ('_TourOptionsMixin', 'Shared options-popup logic for TourDialog and TourToolbar. '
     'Capture destination (Clipboard/File), file mode (save dialog / default location), '
     'directory selector, include adjacent checkbox. Popup-safe directory picker.'),
    ('_JinjaSyntaxHighlighter', 'Regex-based live syntax coloring for Jinja2 + Markdown. '
     'Dark/light theme palette. Error paragraph highlighting (pink background).'),
    ('_ContentViewer', 'Dual-mode widget: view mode (read-only HTML) and edit mode '
     '(monospace editor with syntax highlighting, live validation, Ctrl+wheel zoom, '
     'undo/redo, word wrap). Footer shows validation status + cursor position.'),
    ('TourDialog (DockWindow)', 'Full dock window. Controls: [Options][Prev][List/Progress]'
     '[Next][Capture][Edit] ↔ [Switch to toolbar]. Steps list popup with sorting. '
     'Lifecycle hooks for signal pause/resume on hide/show.'),
    ('TourToolbar (ToolBar)', 'Compact toolbar. Orientation-aware layout (vertical: compact '
     'step display, horizontal: progress bar). Navigation buttons disabled at boundaries. '
     'Expand button opens TourDialog.'),
]

table = doc.add_table(rows=len(classes_dialog) + 1, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Class', 'Description']):
    cell = table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    shade(cell)
for r_idx, (name, desc) in enumerate(classes_dialog):
    table.rows[r_idx + 1].cells[0].text = name
    table.rows[r_idx + 1].cells[0].paragraphs[0].runs[0].bold = True
    table.rows[r_idx + 1].cells[1].text = desc

# 3.3
doc.add_heading('3.3 HM Layer: hmmodeltour.py', level=2)

classes_hm = [
    ('HMModelTour (Singleton)', 'Extends ModelTour. Connects to hmcoresignalspy '
     'signals: BeforeReadModel, AfterReadModel, onModelChange, '
     'ElemsRenumberedInCollector. Clears tour state on model changes. '
     'select_entities() delegates to collect_entities_by_panel().'),
    ('HMDefaultHandler', 'Default fallback for any entity type. Auto-registered for '
     'discovered types (props, mats, loadcols, sets, laminates, plies). '
     'Generic isolate + window_entitymark behavior.'),
    ('HMComponentHandler', 'Handles Components. Isolates with adjacent entity inclusion. '
     'Template: Name, ID, Property, Material, Element count/range, Mass, Thickness.'),
    ('HMPartHandler', 'Handles Parts using name-based identifiers. Template: UID, '
     'Revision, Representation, Components list, Materials, Properties, '
     'Element count/range.'),
    ('HMConnectorHandler', 'Handles Connectors with CE_ReviewConnectors + CE_ReviewLinks '
     'for linked entity isolation. Template: Type, Control, Realization Status, '
     'Layers, Links, Error report, Linked entities.'),
]

table = doc.add_table(rows=len(classes_hm) + 1, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Class', 'Description']):
    cell = table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    shade(cell)
for r_idx, (name, desc) in enumerate(classes_hm):
    table.rows[r_idx + 1].cells[0].text = name
    table.rows[r_idx + 1].cells[0].paragraphs[0].runs[0].bold = True
    table.rows[r_idx + 1].cells[1].text = desc

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Key Functions:').bold = True
funcs = [
    '_register_default_entities() — Auto-discovers entity types via Tcl typeID resolution',
    'collect_entities_by_panel() — Interactive entity selection via *createmarkpanel Tcl',
    '_get_action_location() — Positions panel next to toolbar icon (MVSHM-295089)',
    'run_model_tour() — Entry point: creates TourController and starts the tour',
    '_get_entity_tcl_name_map() — Builds entity class → Tcl name lookup from handlers',
]
for f in funcs:
    doc.add_paragraph(f, style='List Bullet')

# 3.4
doc.add_heading('3.4 HM Layer: hm_utils.py', level=2)

utils_table = doc.add_table(rows=7, cols=2, style='Light Grid Accent 1')
utils_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Function', 'Purpose']):
    cell = utils_table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    shade(cell)

utils_data = [
    ('get_hm_value(entity, attr)', 'Safe attribute access for HM entities'),
    ('hm_getentity(type, id)', 'Fetch an entity by type name + ID'),
    ('hm_getentity_mass(id, type)', 'Get mass via hm_getmass. Formatted output (fixed or scientific)'),
    ('hm_get_comps_elems_details(ids)', 'Element IDs per component(s) via Tcl *createmark. Returns {element_ids, min, max}'),
    ('hm_get_comps_elems(ids)', 'Compatibility wrapper → flat element ID list'),
    ('hm_eval_tcl(command)', 'Safe Tcl eval wrapper for templates'),
]
for r_idx, (name, purpose) in enumerate(utils_data):
    utils_table.rows[r_idx + 1].cells[0].text = name
    utils_table.rows[r_idx + 1].cells[0].paragraphs[0].runs[0].bold = True
    utils_table.rows[r_idx + 1].cells[1].text = purpose

# ══════════════════════════════════════════════════════════════════════════
# SECTION 4
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('4. Complete Workflow', level=1)

doc.add_heading('4.1 Start → Tour Running', level=2)
doc.add_paragraph('The following sequence is executed when run_model_tour() is called:')
steps = [
    '1. run_model_tour(entity_type) calls TourController.start(HMModelTour(), entity_type, mode)',
    '2. tour.select_entities() → HMModelTour.select_entities() → collect_entities_by_panel()',
    '   a. _get_action_location() → determines toolbar bounding box for panel placement',
    '   b. Builds filtered Tcl entity list from registered types present in model',
    '   c. Executes *createmarkpanel Tcl script positioned at toolbar location',
    '   d. Parses result (type name + entity IDs) and builds MDI Collection',
    '3. tour.collection = Collection → triggers current = entities[0]',
    '   a. _set_current() → _on_current_changed() → _notify_observers(EVENT_CURRENT_CHANGED)',
    '4. switch_mode() → shows TourDialog or TourToolbar',
    '5. _activate_tour() registers observer, resumes signals, syncs options, renders content, updates navigation buttons, triggers _review_current() with signal suppression',
]
for s in steps:
    doc.add_paragraph(s)

doc.add_heading('4.2 Navigation Flow', level=2)
p = doc.add_paragraph()
p.add_run('next()').bold = True
p.add_run(' and ')
p.add_run('previous()').bold = True
p.add_run(' both:')
nav_steps = [
    'Get current index via _get_valid_index() (compares by entity ID)',
    'If no valid current → _find_and_set_valid_entity() (forward or reverse)',
    'If at boundary → fire EVENT_REACHED_END / EVENT_REACHED_BEGINNING, wrap to other end',
    'Otherwise → search forward/backward from current ± 1',
    '_review_current() isolates entity with signal suppression',
]
for s in nav_steps:
    doc.add_paragraph(s, style='List Bullet')

p = doc.add_paragraph()
p.add_run('_find_and_set_valid_entity(start_index, reverse):').bold = True
vs_steps = [
    'Forward: range(start_index, len) → first entity where str(entity) succeeds',
    'Reverse: range(start_index, -1, -1) → first valid entity backward',
    'If not found in range → wraps to the other end',
    'If no valid entities at all → _set_current(None)',
]
for s in vs_steps:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('4.3 Event Handling', level=2)

events_table = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
events_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Event', 'Handling']):
    cell = events_table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    shade(cell)
events_data = [
    ('MODEL_DELETE_EVENT', 'Clear collection + current, notify model_deleted=True'),
    ('ENTITY_DELETE_EVENT', '_handle_entity_deletion(): check current validity, find next if deleted'),
    ('ENTITY_UPDATE_EVENT', 'Compare UIDs, notify EVENT_ENTITY_UPDATED if matches current'),
    ('ENTITY_UIDCHANGE_EVENT', 'Resolve old→new UID, re-key captured images, update reference'),
]
for r_idx, (evt, handling) in enumerate(events_data):
    events_table.rows[r_idx + 1].cells[0].text = evt
    events_table.rows[r_idx + 1].cells[1].text = handling

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Signal Suppression:').bold = True
doc.add_paragraph(
    'The _suppressing_signals flag prevents re-entrant ENTITY_UPDATE_EVENT callbacks '
    'during entity isolation. It is set True before handle_entity() and restored False '
    'in the finally block. This flag is checked first thing in _on_entity_deleted() — '
    'before any C++ calls — so bulk SHI from isolateonlyentitybymark costs only a single '
    'Python bool read.'
)

doc.add_heading('4.4 Image Capture Workflow', level=2)
cap_steps = [
    'User clicks Capture button → _resolve_capture_request() determines destination',
    'Clipboard: uses temp PNG file → PIL scale to 900px width → base64 data URI → uiClipboard.SetPixmap',
    'File: prompts save dialog or uses default directory → timestamped PNG (model_tour_YYYY_MM_DD_HH_MM_SS.png)',
    'Store img_tag in _captured_images[entity_uid]',
    'Re-render content with tour.capture() → render_current_with_image() inserts captured image',
]
for s in cap_steps:
    doc.add_paragraph(s, style='List Bullet')

# ══════════════════════════════════════════════════════════════════════════
# SECTION 5
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('5. Template System', level=1)

doc.add_heading('Pipeline', level=2)
pipe = [
    'Jinja2 source (Markdown + Jinja2 syntax)',
    '→ Jinja2 Environment.compile() (with _ImportModuleExtension for {% import_module %})',
    '→ markdown.markdown(result, extensions=[\'extra\', \'smarty\', \'tables\'])',
    '→ _ContentViewer._wrap_html() → CSS-themed full HTML document',
]
for s in pipe:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('Context Available in Templates', level=2)
ctx_table = doc.add_table(rows=3, cols=2, style='Light Grid Accent 1')
ctx_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Source', 'Variables/Functions']):
    cell = ctx_table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    shade(cell)
ctx_data = [
    ('Core utils', 'entity, entity_handler, get_property_value, get_material_info, get_component_children, format_entity_summary, get_entity_link, capture_image, eval_tcl, eval_py'),
    ('HM utils', 'get_value, hm_getentity, hm_getentity_mass, hm_get_comps_elems, hm_get_comps_elems_details, hm module'),
]
for r_idx, (source, vars_) in enumerate(ctx_data):
    ctx_table.rows[r_idx + 1].cells[0].text = source
    ctx_table.rows[r_idx + 1].cells[1].text = vars_

doc.add_heading('Entity-Specific Templates', level=2)
p = doc.add_paragraph()
p.add_run('Component Template:').bold = True
doc.add_paragraph(
    'Displays Name, ID, Property (with name + ID), Material (with name + ID), '
    'Number of Elements, Element ID Range (red if min < 1), Mass, Thickness, '
    'captured image. Uses hm_get_comps_elems_details for bulk element querying.'
)
p = doc.add_paragraph()
p.add_run('Part Template:').bold = True
doc.add_paragraph(
    'Displays Name, ID, UID, Revision (major.study.library), Representation, '
    'Components list (de-duplicated by ID), Materials list, Properties with '
    'Thickness, Number of Elements, Element ID Range, captured image. '
    'All element data fetched in a single bulk call.'
)
p = doc.add_paragraph()
p.add_run('Connector Template:').bold = True
doc.add_paragraph(
    'Displays ID, Type (ce_style), Control ID, Realization Status, Layers, '
    'Link Numbers, Linked Entities grouped by type, Error report, captured image.'
)

doc.add_heading('Persistence', level=2)
persistence = [
    'User edits saved to modeltour_templates.j2 in application settings directory',
    'Each entity type stored as {% block EntityName %}...{% endblock %}',
    'Only modified (non-default) templates are persisted',
    'Templates reverted to default are automatically removed from the file',
    'If no modified templates remain, the file is deleted entirely',
]
for s in persistence:
    doc.add_paragraph(s, style='List Bullet')

# ══════════════════════════════════════════════════════════════════════════
# SECTION 6
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('6. File Inventory & Summary', level=1)

files_table = doc.add_table(rows=6, cols=4, style='Light Grid Accent 1')
files_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['File', 'Lines', 'Layer', 'Key Contents']):
    cell = files_table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    shade(cell)
files_data = [
    ('tour.py', '~2058', 'Core', 'ModelTour (ABC), _Template, _TemplateManager, IEntityHandler, navigation, events, image capture'),
    ('dialog.py', '~2778', 'Core', 'TourController, TourDialog, TourToolbar, _ContentViewer, _TourOptionsMixin, _JinjaSyntaxHighlighter, _TourSettings'),
    ('utils.py', '~265', 'Core', 'Base template helpers: capture_image, eval_tcl/eval_py, property/material access'),
    ('hmmodeltour.py', '~860', 'HM', 'HMModelTour, entity handlers (4), panel selection, auto-discovery, entry point'),
    ('hm_utils.py', '~195', 'HM', 'HM template functions: mass, element details, entity fetching'),
]
for r_idx, (fname, lines, layer, contents) in enumerate(files_data):
    files_table.rows[r_idx + 1].cells[0].text = fname
    files_table.rows[r_idx + 1].cells[1].text = lines
    files_table.rows[r_idx + 1].cells[2].text = layer
    files_table.rows[r_idx + 1].cells[3].text = contents

# ══════════════════════════════════════════════════════════════════════════
# SECTION 7
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('7. Design Patterns', level=1)

patterns_table = doc.add_table(rows=9, cols=2, style='Light Grid Accent 1')
patterns_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Pattern', 'Usage']):
    cell = patterns_table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    shade(cell)
patterns_data = [
    ('Singleton', 'HMModelTour, _TemplateManager, TourController, _TourSettings'),
    ('Abstract Base', 'ModelTour(ABC) with @abstractmethod select_entities()'),
    ('Strategy', 'IEntityHandler per entity type, registered with entity_type= in class statement'),
    ('Observer', 'add_observer / _notify_observers for tour events'),
    ('Mixin', '_TourOptionsMixin shared by TourDialog and TourToolbar'),
    ('Template Method', '_on_current_changed() hook for derived class behavior'),
    ('Lazy Initialization', 'Steps popup created on first show'),
    ('Guard Flag', '_suppressing_signals for re-entrant event prevention, _in_validation for recursive template validation'),
]
for r_idx, (name, usage) in enumerate(patterns_data):
    patterns_table.rows[r_idx + 1].cells[0].text = name
    patterns_table.rows[r_idx + 1].cells[0].paragraphs[0].runs[0].bold = True
    patterns_table.rows[r_idx + 1].cells[1].text = usage

# ══════════════════════════════════════════════════════════════════════════
# SECTION 8
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading('8. Recent Changes (Uncommitted)', level=1)

changes_table = doc.add_table(rows=5, cols=3, style='Light Grid Accent 1')
changes_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['File', 'Change', 'Impact']):
    cell = changes_table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True
    shade(cell)
changes_data = [
    ('hm_utils.py', 'Refactored hm_get_comps_elems → hm_get_comps_elems_details(). Supports multiple component IDs in one Tcl call. Returns {element_ids, min, max} dict. Old function kept as compatibility wrapper.', 'Performance improvement: single Tcl call instead of N calls for N components'),
    ('hmmodeltour.py', 'Updated Component & Part templates to use new bulk API. Registered hm_get_comps_elems_details as template context. Element count changed from entity.elements | length to bulk result length.', 'Faster rendering for parts/components with many elements'),
    ('dialog.py', 'SetMaximumWidth(dpi(300)) on capture destination combo and directory selector. Reduced options popup SetMinimumWidth from 400→300.', 'Prevents oversized widgets in options popup'),
    ('tour.py', 'Fixed backward navigation wrap: removed explicit start_index=self.count-1 in previous(). Added debug print for entity UID in update events.', 'Backward wrap now searches full range, fixing potential infinite loops at beginning of collection'),
]
for r_idx, (fname, change, impact) in enumerate(changes_data):
    changes_table.rows[r_idx + 1].cells[0].text = fname
    changes_table.rows[r_idx + 1].cells[1].text = change
    changes_table.rows[r_idx + 1].cells[2].text = impact

# ══════════════════════════════════════════════════════════════════════════
# FOOTER
# ══════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— End of Document —')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.size = Pt(10)

# ── Save ────────────────────────────────────────────────────────────────
output_path = os.path.join(os.path.dirname(__file__), 'ModelTour_Documentation.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
